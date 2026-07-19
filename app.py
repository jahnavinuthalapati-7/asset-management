import streamlit as st
from supabase import create_client, Client 
import hashlib
from datetime import date, datetime
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import io  
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Alignment, Border, Side , Font
from docx import Document
from docx.shared import Inches, Pt, RGBColor  
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============ SUPABASE CONFIG - ADD YOUR KEYS HERE ============
SUPABASE_URL = "https://uthdmnqizwslgpxrbork.supabase.co"  
SUPABASE_KEY ="eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InV0aGRtbnFpendzbGdweHJib3JrIiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODQyNjc5NTAsImV4cCI6MjA5OTg0Mzk1MH0.gdV3duAKYgLKWU5R6gUOHugdf87kWKTVUcVZ2CLkMvs"
# ============================================================

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Constants  
ASSET_TYPES = ["Laptop", "i Pad", "Desktop", "Phone"] 

LOCATIONS = [
    "Bird Automotive Pvt.Ltd 4IDC",
    "Bird Automotive Pvt.Ltd  Baani",
    "Bird Automotive Pvt.Ltd Service Factory ",
    "Bird Automotive Pvt.Ltd Mini Delhi",
    "Bird Automotive Pvt.Ltd Dehradun",
    "Bird Automobile Porsche Kolkata WS",
    "Bird Automobile Porsche Kolkata SR",
    "Bird Premium Selection Bentley SR",
    "Bird Premium Selection Bentley WS",
    "E-Next Mobility Citroen 48",
    "E-Next Mobility Sector-29 GGN",
]

def hash_pw(password):
    return hashlib.sha256(password.encode()).hexdigest()

# ============ DATABASE OPERATIONS ============

def add_user(username, password, role):
    try:
        response = supabase.table("users").insert({
            "username": username,
            "password_hash": hash_pw(password),
            "role": role
        }).execute()
        return True
    except Exception as e:
        return False

def verify_user(username, password):
    try:
        response = supabase.table("users").select("*").eq("username", username).execute()
        if response.data:
            user = response.data[0]
            if user["password_hash"] == hash_pw(password):
                return user["role"]
        return None
    except:
        return None

def change_password(username, old_password, new_password):
    try:
        response = supabase.table("users").select("password_hash").eq("username", username).execute()
        if response.data:
            user = response.data[0]
            if user["password_hash"] == hash_pw(old_password):
                supabase.table("users").update({"password_hash": hash_pw(new_password)}).eq("username", username).execute()
                return True
        return False
    except:
        return False

def add_asset(asset_type, model_name, serial_number, vendor_name, purchase_date, department, 
              location, designation, current_holder, employee_code, employee_mobile, date_issued, remarks):
    try:
        response = supabase.table("laptops").insert({
            "asset_type": asset_type,
            "model_name": model_name,
            "serial_number": serial_number,
            "vendor_name": vendor_name,
            "purchase_date": purchase_date,
            "department": department,
            "location": location,
            "designation": designation,
            "current_holder": current_holder,
            "employee_code": employee_code,
            "employee_mobile": employee_mobile,
            "date_issued": date_issued,
            "status": "Active",
            "remarks": remarks
        }).execute()
        
        if response.data:
            laptop_id = response.data[0]["id"]
            supabase.table("history").insert({
                "laptop_id": laptop_id,
                "person_name": current_holder,
                "department": department,
                "location": location,
                "designation": designation,
                "date_issued": date_issued,
                "remarks": remarks
            }).execute()
        return True
    except Exception as e:
        st.error(f"Error: {str(e)}")
        return False

def get_all_assets():
    try:
        response = supabase.table("laptops").select("*").execute()
        return response.data if response.data else []
    except:
        return []

def get_asset_by_id(asset_id):
    try:
        response = supabase.table("laptops").select("*").eq("id", asset_id).execute()
        return response.data[0] if response.data else None
    except:
        return None

def get_asset_history(asset_id):
    try:
        response = supabase.table("history").select("*").eq("laptop_id", asset_id).order("created_at", desc=True).execute()
        return response.data if response.data else []
    except:
        return []

def reallocate_asset(asset_id, new_holder, new_department, new_location, new_designation, reallocation_date, remarks):
    try:
        supabase.table("laptops").update({
            "current_holder": new_holder,
            "department": new_department,
            "location": new_location,
            "designation": new_designation,
            "employee_code": new_holder,
            "reallocation_date": str(reallocation_date)
        }).eq("id", asset_id).execute()
        
        supabase.table("history").insert({
            "laptop_id": asset_id,
            "person_name": new_holder,
            "department": new_department,
            "location": new_location,
            "designation": new_designation,
            "date_issued": str(reallocation_date),
            "remarks": remarks
        }).execute()
        return True
    except:
        return False

def mark_as_scrap(asset_id):
    try:
        supabase.table("laptops").update({
            "status": "Scrap",
            "current_holder": None
        }).eq("id", asset_id).execute()
        return True
    except:
        return False

def delete_asset(asset_id):
    try:
        supabase.table("history").delete().eq("laptop_id", asset_id).execute()
        supabase.table("laptops").delete().eq("id", asset_id).execute()
        return True
    except:
        return False

# ============ EXPORT FUNCTIONS ============

def export_to_pdf(filtered_data, columns):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    
    title_style = getSampleStyleSheet()['Heading1']
    elements.append(Paragraph("Asset Management Report", title_style))
    elements.append(Spacer(1, 12))
    
    data = [[col.upper() for col in columns]]
    for row in filtered_data:
        data.append([str(val) if val else "-" for val in row])
    
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_excel(filtered_data, columns):
    wb = Workbook()
    ws = wb.active
    ws.title = "Assets"
    
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_obj = Font(bold=True, color="FFFFFF")
    
    ws.append(columns)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font_obj
    
    for row in filtered_data:
        ws.append(row)
    
    xlsx_buffer = io.BytesIO()
    wb.save(xlsx_buffer)
    xlsx_buffer.seek(0)
    return xlsx_buffer.getvalue() 

def export_to_word(filtered_data, columns):
    doc = Document()
    doc.add_heading("Asset Management Report", 0)
    
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col.upper()
    
    for row in filtered_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val) if val else "-"
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()

# ============ UI FUNCTIONS ============

def apply_theme():
    st.markdown("""
        <style>
        .stApp {
            background: linear-gradient(160deg, #dbe9ff 0%, #eef2ff 45%, #fef6e8 100%);
        }
        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #2b3a67 0%, #1f2a4a 100%);
        }
        section[data-testid="stSidebar"] * {
            color: #f5f7fa !important;
        }
        h1, h2, h3 {
            color: #1f2a4a;
        }
        div.stButton > button {
            background-color: #3b6ef0;
            color: white;
            border-radius: 8px;
            border: none;
        }
        div.stButton > button:hover {
            background-color: #2b53c4;
        }
        </style>
    """, unsafe_allow_html=True)

def show_login():
    st.set_page_config(page_title="Asset Management", layout="wide", initial_sidebar_state="collapsed")
    apply_theme()
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        st.markdown("#  Asset Management")
        st.markdown("---")
        
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        
        if st.button("Login", use_container_width=True):
            if username and password:
                authenticated_role = verify_user(username, password)
                if authenticated_role:
                    st.session_state["logged_in"] = True
                    st.session_state["username"] = username
                    st.session_state["role"] = authenticated_role
                    st.rerun()
                else:
                    st.error("Invalid credentials")
            else:
                st.error("Please enter username and password")

def show_dashboard(role):
    st.header(" Asset Inventory")
    
    all_assets = get_all_assets()
    
    if not all_assets:
        st.info("No assets found. Go to 'Add Asset' to add your first asset.")
        return
    
    search_query = st.text_input(" Search (model name, serial number, department, holder name, vendor name, issued date or location)")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        departments = sorted(set([a.get("department", "") for a in all_assets if a.get("department")]))
        dept_filter = st.selectbox("Filter by Department", ["All"] + departments)
    with col2:
        locations = sorted(set([a.get("location", "") for a in all_assets if a.get("location")]))
        loc_filter = st.selectbox("Filter by Location", ["All"] + locations)
    with col3:
        types = sorted(set([a.get("asset_type", "") for a in all_assets if a.get("asset_type")]))
        type_filter = st.selectbox("Filter by Asset Type", ["All"] + types)
    
    filtered = all_assets
    
    if dept_filter != "All":
        filtered = [a for a in filtered if a.get("department") == dept_filter]
    if loc_filter != "All":
        filtered = [a for a in filtered if a.get("location") == loc_filter]
    if type_filter != "All":
        filtered = [a for a in filtered if a.get("asset_type") == type_filter]
    
    if search_query.strip():
        q = search_query.strip().lower()
        search_fields = ["model_name", "serial_number", "department", "current_holder", "vendor_name", "employee_code", "location"]
        filtered = [
            a for a in filtered
            if any(q in str(a.get(field, "")).lower() for field in search_fields)
        ]
    
    if filtered:
        for asset in filtered:
            label = f" [{asset.get('asset_type', 'Asset')}] {asset['model_name']} — {asset.get('current_holder', 'Unassigned')}"
            with st.expander(label):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Asset Type:** {asset.get('asset_type')}")
                    st.write(f"**Serial Number:** {asset.get('serial_number')}")
                    st.write(f"**Vendor:** {asset.get('vendor_name')}")
                    st.write(f"**Purchase Date:** {asset.get('purchase_date')}")
                    st.write(f"**Department:** {asset.get('department')}")
                    st.write(f"**Location:** {asset.get('location')}")
                
                with col2:
                    st.write(f"**Designation:** {asset.get('designation')}")
                    st.write(f"**Current Holder:** {asset.get('current_holder')}")
                    st.write(f"**Employee Code:** {asset.get('employee_code')}")
                    st.write(f"**Employee Mobile:** {asset.get('employee_mobile')}")
                    st.write(f"**Date Issued:** {asset.get('date_issued')}")
                    st.write(f"**Status:** {asset.get('status')}")
                
                if role == "head_admin":
                    st.write(f"**Reallocation Date:** {asset.get('reallocation_date', '-')}")
                    st.markdown("** Assignment History**")
                    history = get_asset_history(asset['id'])
                    if history:
                        for h in history:
                            st.write(f"- {h['person_name']} ({h['designation']}, {h['department']}) — {h['date_issued']}")
                
                if asset.get("remarks"):
                    st.write(f"**Remarks:** {asset['remarks']}")
                
                if role == "head_admin":
                    if st.button(" Delete", key=f"del_{asset['id']}"):
                        delete_asset(asset['id'])
                        st.success("Deleted!")
                        st.rerun()
    else:
        st.warning("No assets match your filters")

def show_add_asset():
    st.header(" Add New Asset")
    with st.form("add_asset_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            asset_type = st.selectbox("Asset Type", ASSET_TYPES)
            model_name = st.text_input("Model Name")
            serial_number = st.text_input("Asset Serial Number")
            vendor_name = st.text_input("Vendor Name")
            purchase_date = st.date_input("Purchase Date")
        
        with col2:
            department = st.text_input("Department")
            location = st.selectbox("Location", LOCATIONS)
            designation = st.text_input("Designation")
            current_holder = st.text_input("Name of Person")
            employee_code = st.text_input("Employee Code")
        
        employee_mobile = st.text_input("Employee Mobile Number")
        date_issued = st.date_input("Date Issued", value=date.today())
        remarks = st.text_area("Remarks (optional)")
        
        submitted = st.form_submit_button("Add Asset")
        if submitted:
            if not model_name:
                st.error("Model name is required.")
            else:
                if add_asset(asset_type, model_name, serial_number, vendor_name, str(purchase_date),
                            department, location, designation, current_holder, employee_code,
                            employee_mobile, str(date_issued), remarks):
                    st.success(f"Asset '{model_name}' added successfully!")
                else:
                    st.error("Failed to add asset")

def show_reallocate():
    st.header(" Reallocate Asset")
    
    all_assets = get_all_assets()
    if not all_assets:
        st.info("No assets available to reallocate")
        return
    
    asset_options = {f"{a['id']} - {a['model_name']}": a for a in all_assets}
    selected_label = st.selectbox("Select Asset", asset_options.keys())
    asset = asset_options[selected_label]
    current_location = asset.get("location")
    
    st.info(f"Current Holder: {asset.get('current_holder', 'N/A')}")
    
    asset_action = st.radio("What do you want to do?", ["Reallocate to Person", "Mark as Scrap"])
    
    if asset_action == "Mark as Scrap":
        if st.button("Mark as Scrap"):
            if mark_as_scrap(asset['id']):
                st.success("Asset marked as Scrap!")
                st.rerun()
    else:
        with st.form("reallocate_form"):
            new_holder = st.text_input("New Holder's Name")
            new_department = st.text_input("New Department")
            new_employee_code = st.text_input("New Employee Code")
            new_employee_mobile = st.text_input("New Employee Mobile")
            
            location_options = LOCATIONS.copy()
            if current_location and current_location not in location_options:
                location_options.append(current_location)
            default_idx = location_options.index(current_location) if current_location in location_options else 0
            new_location = st.selectbox("New Location", location_options, index=default_idx)
            
            new_designation = st.text_input("New Designation")
            reallocation_date = st.date_input("Reallocation Date", value=date.today())
            remarks = st.text_area("Remarks (optional)")
            
            submitted = st.form_submit_button("Reallocate")
            if submitted:
                if not new_holder:
                    st.error("Holder name required")
                else:
                    if reallocate_asset(asset['id'], new_holder, new_department, new_location, new_designation, reallocation_date, remarks):
                        st.success("Asset reallocated!")
                        st.rerun()

def show_change_password():
    st.header(" Change My Password")
    username = st.session_state["username"]
    
    with st.form("change_password_form"):
        old_password = st.text_input("Current Password", type="password")
        new_password = st.text_input("New Password", type="password")
        confirm_password = st.text_input("Confirm New Password", type="password")
        
        submitted = st.form_submit_button("Update Password")
        if submitted:
            if not old_password or not new_password:
                st.error("All fields required")
            elif new_password != confirm_password:
                st.error("Passwords don't match")
            else:
                if change_password(username, old_password, new_password):
                    st.success("Password updated!")
                else:
                    st.error("Current password incorrect")

def show_reports(role):
    st.header(" Asset Reports")
    
    all_assets = get_all_assets()
    if not all_assets:
        st.info("No assets to report")
        return
    
    search_query = st.text_input("🔍 Search by vendor, holder, department, location, serial number, or date")
    
    col1, col2 = st.columns(2)
    with col1:
        status_filter = st.selectbox("Filter by Status", ["All", "Active", "Scrap", "Inactive"])
    with col2:
        export_format = st.selectbox("Export As", ["PDF", "Excel", "Word"])
    
    filtered = all_assets
    
    if status_filter != "All":
        filtered = [a for a in filtered if a.get("status") == status_filter]
    
    if search_query.strip():
        q = search_query.strip().lower()
        filtered = [
            a for a in filtered
            if any(q in str(a.get(field, "")).lower() for field in 
                  ["vendor_name", "current_holder", "department", "location", "serial_number", "employee_code", "status"])
        ]
    
    if filtered:
        st.write(f"**Total Assets: {len(filtered)}**")
        
        df_data = []
        for a in filtered:
            row = {
                "ID": a["id"],
                "Type": a.get("asset_type"),
                "Model": a.get("model_name"),
                "Serial": a.get("serial_number"),
                "Vendor": a.get("vendor_name"),
                "Emp Code": a.get("employee_code"),
                "Emp Mobile": a.get("employee_mobile"),
                "Department": a.get("department"),
                "Location": a.get("location"),
                "Holder": a.get("current_holder"),
                "Status": a.get("status"),
                "Issued": a.get("date_issued")
            }
            if role == "head_admin":
                row["Reallocated"] = a.get("reallocation_date", "-")
            df_data.append(row)
        
        st.dataframe(df_data, use_container_width=True)
        
        st.markdown("---")
        
        if export_format == "PDF":
            export_data = [[str(v) for v in d.values()] for d in df_data]
            export_cols = list(df_data[0].keys())
            pdf_buffer = export_to_pdf(export_data, export_cols)
            st.download_button(" Download PDF", pdf_buffer, "report.pdf", "application/pdf")
        
        elif export_format == "Excel":
            export_data = [[str(v) for v in d.values()] for d in df_data]
            export_cols = list(df_data[0].keys())
            excel_buffer = export_to_excel(export_data, export_cols)
            st.download_button(" Download Excel", excel_buffer, "report.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        
        elif export_format == "Word":
            export_data = [[str(v) for v in d.values()] for d in df_data]
            export_cols = list(df_data[0].keys())
            word_buffer = export_to_word(export_data, export_cols)
            st.download_button(" Download Word", word_buffer, "report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("No assets match filters")

def show_manage_users():
    st.header(" Manage Users")
    
    with st.form("add_user_form"):
        new_username = st.text_input("New Username")
        new_password = st.text_input("New Password", type="password")
        new_role = st.selectbox("Role", ["admin", "head_admin"])
        
        submitted = st.form_submit_button("Create User")
        if submitted:
            if not new_username or not new_password:
                st.error("Username and password required")
            else:
                if add_user(new_username, new_password, new_role):
                    st.success(f"User '{new_username}' created!")
                else:
                    st.error("Username already exists")

# ============ MAIN APP ============

def main():
    st.set_page_config(page_title="Asset Management", layout="wide")
    apply_theme()
    
    if "logged_in" not in st.session_state:
        st.session_state["logged_in"] = False
    
    if not st.session_state["logged_in"]:
        show_login()
    else:
        st.sidebar.markdown(f"## Welcome, {st.session_state['username']}!")
        
        pages = ["Dashboard", "Add Asset", "Reallocate Asset", "Reports", "Change Password"]
        if st.session_state["role"] == "head_admin":
            pages.append("Manage Users")
        
        page = st.sidebar.radio("Navigate", pages)
        
        if st.sidebar.button("Logout"):
            st.session_state["logged_in"] = False
            st.rerun()
        
        st.title(" Asset Management")
        
        if page == "Dashboard":
            show_dashboard(st.session_state["role"])
        elif page == "Add Asset":
            show_add_asset()
        elif page == "Reallocate Asset":
            show_reallocate()
        elif page == "Reports":
            show_reports(st.session_state["role"])
        elif page == "Change Password":
            show_change_password()
        elif page == "Manage Users":
            show_manage_users()

if __name__ == "__main__":
    main()
